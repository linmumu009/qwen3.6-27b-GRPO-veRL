#!/usr/bin/env python3
"""Build the Qwen3.8 reasoning-effort investigation DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/qwen38_reasoning_effort_investigation_20260817.md"
OUTPUT = ROOT / "docs/qwen38_reasoning_effort_investigation_20260817.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 34, 52)
MUTED = RGBColor(95, 105, 120)
LIGHT_FILL = "F2F4F7"
CODE_FILL = "F6F8FA"
CAUTION_FILL = "FFF7E0"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_rich_text(paragraph, text, *, size=11, color=INK, bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=size - 0.5,
                         color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles["Code Block"]
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.3,
                     color=INK)


def add_callout(doc, label, text, fill=CAUTION_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}：")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    add_rich_text(p, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, str(value), size=9.2)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def parse_table(lines, start):
    headers = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    rows = []
    idx = start + 2
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        rows.append([c.strip() for c in lines[idx].strip().strip("|").split("|")])
        idx += 1
    return headers, rows, idx


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9.3)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.05
    ppr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    ppr.append(shd)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Qwen3.8 Reasoning Effort | 技术调研")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("QWEN3.8 REASONING EFFORT")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("实现机制、训练范式与Max/27B对比")
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("0号机实机代码审计与千问官方资料核对")
    set_run_font(run, size=12.5, color=MUTED)

    for label, value in (
        ("调研对象", "Qwen3.8-27B / PI 0.82.1 / vLLM 0.23.0 / Qwen3.8-Max"),
        ("调研日期", "2026-08-17"),
        ("证据口径", "实机文件与运行代码优先；官方公开资料其次；推断与事实分离"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}：")
        set_run_font(run, size=10.5, color=INK, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(12)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pbdr.append(bottom)
    ppr.append(pbdr)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    i = 6  # Skip source title and metadata already rendered in masthead.
    in_code = False
    code_lines = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines or [""])
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(":", "").replace("-", "").strip()) == set():
            headers, rows, next_i = parse_table(lines, i)
            if len(headers) == 2:
                widths = [2700, 6660]
            elif len(headers) == 3:
                widths = [2100, 3300, 3960]
            else:
                widths = [1800, 1890, 1890, 1890, 1890][:len(headers)]
                if sum(widths) != 9360:
                    widths[-1] += 9360 - sum(widths)
            add_table(doc, headers, rows, widths)
            i = next_i
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("**") and "：**" in stripped:
            label, text = stripped.split("：**", 1)
            add_callout(doc, label.strip("*"), text.strip())
            i += 1
            continue
        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    props = doc.core_properties
    props.title = "Qwen3.8 Reasoning Effort：实现机制、训练范式与Max/27B对比"
    props.subject = "Qwen3.8-27B reasoning_effort 实机调研"
    props.author = "LLIN Research"
    props.keywords = "Qwen3.8, reasoning_effort, thinking, training, vLLM"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
